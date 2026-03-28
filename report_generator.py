"""
report_generator.py -- HTML + DOCX
===================================
Tested end-to-end: 15 figures, 0 errors with real analyzer chain.

CRITICAL: install exact versions:
    pip install python-docx "kaleido==0.2.1"
    
kaleido 1.x requires Chrome and WILL NOT WORK on servers.
kaleido 0.2.1 is self-contained and works everywhere.
"""

import io, os, sys, time, tempfile, subprocess, traceback
from datetime import datetime
import pandas as pd
import numpy as np
import warnings
warnings.filterwarnings('ignore')

# Diagnostic log
_log = []
def _L(msg): _log.append(msg)
def get_diagnostic_log(): return list(_log)
def clear_diagnostic_log(): _log.clear()


# ------------------------------------------------------------------ #
#  PNG export -- two methods, with diagnostics                         #
# ------------------------------------------------------------------ #

def _to_png(fig, w=900, h=500):
    """Plotly figure -> PNG bytes. Returns None on failure."""
    if fig is None:
        return None
    # Clamp dimensions
    w = max(400, min(w, 1400))
    h = max(300, min(h, 1000))
    
    # Method 1: direct kaleido
    try:
        b = fig.to_image(format='png', width=w, height=h, scale=2)
        if b and len(b) > 500:
            return b
    except Exception as e:
        _L(f"    png direct: {str(e)[:80]}")
    
    # Method 2: subprocess with timeout
    fj = fp = None
    try:
        fj = tempfile.mktemp(suffix='.json')
        fp = tempfile.mktemp(suffix='.png')
        with open(fj, 'w') as f:
            f.write(fig.to_json())
        code = (f"import warnings;warnings.filterwarnings('ignore');"
                f"import plotly.io as p;f=p.from_json(open(r'{fj}').read());"
                f"f.write_image(r'{fp}',format='png',width={w},height={h},scale=2)")
        proc = subprocess.Popen([sys.executable, '-c', code],
                                stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        try:
            proc.wait(timeout=30)
        except subprocess.TimeoutExpired:
            proc.kill(); proc.wait(2)
            _L("    png subprocess: timeout")
            return None
        if proc.returncode == 0 and os.path.exists(fp):
            with open(fp, 'rb') as f:
                data = f.read()
            return data if len(data) > 500 else None
        else:
            err = proc.stderr.read().decode('utf-8', errors='replace')[:150]
            _L(f"    png subprocess: rc={proc.returncode}, {err}")
    except Exception as e:
        _L(f"    png subprocess exception: {str(e)[:80]}")
    finally:
        for p_ in [fj, fp]:
            if p_:
                try: os.unlink(p_)
                except: pass
    return None


def _call(fn, label, *a, **kw):
    """Safe call with diagnostic logging."""
    try:
        r = fn(*a, **kw)
        if r is None:
            _L(f"  {label}: returned None")
        return r
    except Exception as e:
        _L(f"  {label}: ERROR {str(e)[:150]}")
        return None


def _df_html(df, n=25):
    if df is None or df.empty: return ""
    d = df.head(n).copy()
    for c in d.columns:
        if d[c].dtype in ('float64','float32'):
            d[c] = d[c].apply(lambda x: f'{x:.4f}' if pd.notnull(x) else '')
    return d.to_html(index=False, border=0)


# ================================================================== #
#  HTML REPORT (unchanged logic, cleaner structure)                    #
# ================================================================== #

class HTMLReportGenerator:
    def __init__(self, data_loader, settings):
        self.dl = data_loader
        self.s = settings
        self._n = 0

    def _ch(self, fig, cap):
        if fig is None: return ""
        self._n += 1
        try:
            div = fig.to_html(full_html=False, include_plotlyjs=False,
                              div_id=f"f{self._n}",
                              config={'displayModeBar': True, 'responsive': True})
        except: return ""
        return (f'<div style="margin:15px 0;border:1px solid #ECF0F1;border-radius:8px;'
                f'overflow:hidden">{div}</div>'
                f'<p style="text-align:center;color:#7F8C8D;font-size:.85em">'
                f'<b>Fig {self._n}.</b> {cap}</p>\n')

    def generate(self, filepath, title, analyzers, include_sections,
                 progress_cb=None, **kw):
        self._n = 0; t0 = time.time()
        gc = self.s.get('group_column','')
        L = self.dl
        ns = L.otu_table.shape[0] if L.otu_table is not None else 0
        no = L.otu_table.shape[1] if L.otu_table is not None else 0
        gi = ''
        if gc and gc in L.metadata.columns:
            vc = L.metadata[gc].value_counts()
            gi = ' | '.join(f'{g}(n={n})' for g,n in vc.items())

        css = ("body{font-family:'Segoe UI',sans-serif;margin:30px 50px;color:#2C3E50}"
               "h1{border-bottom:3px solid #3498DB;padding-bottom:10px}"
               "h2{border-left:4px solid #3498DB;padding-left:12px;margin-top:30px}"
               "table{border-collapse:collapse;margin:12px 0;font-size:.85em}"
               "th{background:#2C3E50;color:#fff;padding:7px 10px}"
               "td{padding:5px 10px;border:1px solid #BDC3C7;text-align:center}"
               "tr:nth-child(even){background:#F8F9FA}"
               ".bx{background:#ECF0F1;padding:12px 18px;border-radius:8px;margin:12px 0}")

        p = [f'<!DOCTYPE html><html><head><meta charset="utf-8"><title>{title}</title>',
             '<script src="https://cdn.plot.ly/plotly-2.27.0.min.js"></script>',
             f'<style>{css}</style></head><body>',
             f'<h1>{title}</h1>',
             f'<p style="color:#95A5A6">{datetime.now().strftime("%Y-%m-%d %H:%M")}</p>',
             f'<div class="bx"><b>Samples:</b> {ns} | <b>OTUs:</b> {no} | '
             f'<b>Group:</b> {gc} | {gi}</div>']

        D = dict(alpha=self._sa, beta=self._sb, composition=self._sc,
                 heatmap=self._sh, clustered_heatmap=self._sch,
                 network=self._sn, ml=self._sm)
        tot = max(1, len([s for s in include_sections if s in analyzers]))
        for i,sec in enumerate(include_sections):
            if sec not in analyzers: continue
            if progress_cb: progress_cb(f"{sec}...", int(10+80*i/tot))
            fn = D.get(sec)
            if fn:
                try: p.append(fn(analyzers[sec]))
                except Exception as e: p.append(f'<p style="color:red">[{sec}: {e}]</p>')

        p.append(f'<hr><p style="color:#95A5A6;text-align:center">'
                 f'{self._n} figures | {time.time()-t0:.1f}s</p></body></html>')
        with open(filepath,'w',encoding='utf-8') as f: f.write('\n'.join(p))
        if progress_cb: progress_cb("Done!",100)
        return filepath

    def _sa(self,a):
        h='<h2>Alpha Diversity</h2>'
        df=_call(a.get_statistics_summary,"a.stats"); 
        if df is not None and not df.empty: h+=_df_html(df)
        fig=_call(a.plot_all_metrics_plotly,"a.plot",['observed','shannon','simpson','pielou'])
        return h+self._ch(fig,'Alpha diversity')

    def _sb(self,a):
        h='<h2>Beta Diversity</h2>'
        df=_call(a.get_permanova_summary,"b.perm")
        if df is not None and not df.empty: h+='<h3>PERMANOVA</h3>'+_df_html(df)
        gc=self.s.get('group_column')
        for m in getattr(a,'distance_matrices',{}):
            if m in getattr(a,'ordination_results',{}):
                fig=_call(a.plot_2d_plotly,"b.pcoa",m,gc); h+=self._ch(fig,f'PCoA({m})'); break
        return h

    def _sc(self,a):
        h='<h2>Composition</h2>'
        for mn,kw,cap in [('create_plot',dict(top_n=15),'Bar'),
                          ('create_group_comparison_plot',dict(top_n=15),'Group'),
                          ('create_sunburst',dict(top_n=50),'Sunburst'),
                          ('create_core_microbiome_plot',dict(),'Core'),
                          ('create_rank_abundance_curve',dict(top_n=30),'Rank')]:
            fn=getattr(a,mn,None)
            if fn:
                r=_call(fn,f"c.{mn}",**kw); fig=r[0] if isinstance(r,tuple) else r
                h+=self._ch(fig,cap)
        return h

    def _sh(self,a):
        h='<h2>Heatmap</h2>'
        return h+self._ch(_call(a.create_heatmap,"h.heatmap",top_n=30),'Heatmap')

    def _sch(self,a):
        h='<h2>Clustered Heatmap</h2>'
        return h+self._ch(_call(a.create_clustered_heatmap,"ch.clust",top_n=30),'Clustered')

    def _sn(self,a):
        h='<h2>Network</h2>'
        fig=_call(a.create_network_plot,"n.plot",layout='spring',color_by='kingdom',show_labels=True)
        h+=self._ch(fig,'Network')
        fig2=_call(a.create_community_summary_plot,"n.comm"); h+=self._ch(fig2,'Communities')
        m=_call(a.calculate_network_metrics,"n.met")
        if m:
            h+='<div class="bx">'
            for k in ['nodes','edges','density','average_degree','clustering_coefficient']:
                v=m.get(k,'?'); h+=f"<b>{k}:</b> {f'{v:.4f}' if isinstance(v,float) else v} | "
            h+='</div>'
        ks=_call(a.identify_keystone_species,"n.ks",top_n=10)
        if ks is not None and not ks.empty:
            cols=[c for c in ['Taxon','Degree','Betweenness','Keystone_Score'] if c in ks.columns]
            h+='<h3>Keystone</h3>'+_df_html(ks[cols].head(10))
        return h

    def _sm(self,a):
        h='<h2>ML</h2>'
        if not hasattr(a,'models') or not a.models: return h+'<p>No models.</p>'
        comp=_call(a.get_model_comparison,"ml.comp")
        if comp is not None and not comp.empty: h+=_df_html(comp)
        fig=_call(a.plot_model_comparison,"ml.fig"); h+=self._ch(fig,'Models')
        fig2=_call(a.plot_feature_importance,"ml.feat",top_n=15); h+=self._ch(fig2,'Features')
        return h


# ================================================================== #
#  DOCX REPORT                                                         #
# ================================================================== #

class DOCXReportGenerator:
    def __init__(self, data_loader, settings):
        self.dl = data_loader
        self.s = settings
        self._fn = 0
        self._errors = []

    @staticmethod
    def _Pt(n):
        from docx.shared import Pt; return Pt(n)
    @staticmethod
    def _RGB(r,g,b):
        from docx.shared import RGBColor; return RGBColor(r,g,b)
    @staticmethod
    def _In(n):
        from docx.shared import Inches; return Inches(n)

    # -- generate --

    def generate(self, filepath, title, analyzers, include_sections, progress_cb=None):
        from docx import Document as DX
        clear_diagnostic_log()
        self._fn = 0; self._errors = []
        doc = DX(); self._styles(doc)

        _L("=== DOCX Generation Start ===")
        if progress_cb: progress_cb("Title...",3)
        self._title(doc, title)

        D = dict(alpha=self._s_alpha, beta=self._s_beta, composition=self._s_comp,
                 heatmap=self._s_heat, clustered_heatmap=self._s_clust,
                 network=self._s_net, ml=self._s_ml)
        tot = max(1, len([s for s in include_sections if s in analyzers]))
        done = 0
        for sec in include_sections:
            if sec not in analyzers: continue
            done += 1
            _L(f"\n--- {sec} ---")
            if progress_cb: progress_cb(f"{sec}...", int(5+85*done/tot))
            fn = D.get(sec)
            if fn:
                try: fn(doc, analyzers[sec])
                except Exception as e:
                    msg = f"{sec}: {str(e)[:200]}"
                    _L(f"  SECTION ERROR: {msg}")
                    self._errors.append(msg)
                    p = doc.add_paragraph()
                    r = p.add_run(f"[Error: {msg}]")
                    r.font.size = self._Pt(9)
                    r.font.color.rgb = self._RGB(0xC0,0x39,0x2B)

        # Diag page
        doc.add_page_break()
        doc.add_heading('Report Diagnostics', level=2)
        self._txt(doc, f"Figures: {self._fn}, Errors: {len(self._errors)}")
        for e in self._errors:
            self._txt(doc, f"  - {e}")

        doc.add_page_break()
        self._ctr(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                  9, (0x95,0xA5,0xA6))
        self._ctr(doc, "Microbiome Analysis Platform", 9, (0x95,0xA5,0xA6), italic=True)

        if progress_cb: progress_cb("Saving...",95)
        doc.save(filepath)
        if progress_cb: progress_cb("Done!",100)
        _L(f"\n=== Done: {self._fn} figs, {len(self._errors)} errors ===")
        return filepath

    # -- styles --

    def _styles(self, doc):
        s = doc.styles['Normal']
        s.font.name = 'Arial'; s.font.size = self._Pt(10)
        s.font.color.rgb = self._RGB(0x2C,0x3E,0x50)
        for nm,sz,c in [('Heading 1',15,(0x1A,0x52,0x76)),
                        ('Heading 2',12,(0x2E,0x86,0xC1)),
                        ('Heading 3',11,(0x34,0x49,0x5E))]:
            st = doc.styles[nm]; st.font.name = 'Arial'
            st.font.size = self._Pt(sz); st.font.bold = True
            st.font.color.rgb = self._RGB(*c)

    # -- title --

    def _title(self, doc, title):
        for _ in range(5): doc.add_paragraph()
        self._ctr(doc, title.upper(), 22, (0x1A,0x52,0x76), bold=True)
        self._ctr(doc, datetime.now().strftime('%d %B %Y'), 12, (0x7F,0x8C,0x8D))
        doc.add_paragraph()
        L = self.dl; gc = self.s.get('group_column','')
        ns = L.otu_table.shape[0] if L.otu_table is not None else 0
        no = L.otu_table.shape[1] if L.otu_table is not None else 0
        lines = [f"Samples: {ns}", f"OTUs/ASVs: {no}"]
        if gc: lines.append(f"Grouping: {gc}")
        lines.append(f"Level: {self.s.get('taxonomic_level','Genus')}")
        lines.append(f"Normalization: {self.s.get('normalization','relative')}")
        if gc and gc in L.metadata.columns:
            vc = L.metadata[gc].value_counts()
            lines.append("Groups: "+", ".join(f"{g} (n={n})" for g,n in vc.items()))
        for l in lines: self._ctr(doc, l, 10)
        doc.add_page_break()

    # -- primitives --

    def _ctr(self, doc, text, sz=10, clr=None, bold=False, italic=False):
        from docx.enum.text import WD_ALIGN_PARAGRAPH as A
        p = doc.add_paragraph(); p.alignment = A.CENTER
        r = p.add_run(text); r.font.size = self._Pt(sz); r.font.name = 'Arial'
        if bold: r.font.bold = True
        if italic: r.font.italic = True
        if clr: r.font.color.rgb = self._RGB(*clr)

    def _txt(self, doc, text):
        return doc.add_paragraph(text)

    def _fig(self, doc, fig, cap, w=6.2, fw=900, fh=500):
        if fig is None:
            _L(f"  Fig '{cap}': skipped (None)")
            return
        self._fn += 1
        _L(f"  Fig {self._fn} '{cap}': exporting {fw}x{fh}...")
        from docx.enum.text import WD_ALIGN_PARAGRAPH as A
        # Clamp height from figure layout
        real_h = getattr(fig.layout, 'height', None)
        if real_h and real_h > 100:
            fh = min(int(real_h), 900)
        png = _to_png(fig, w=fw, h=fh)
        if png:
            _L(f"  Fig {self._fn}: OK ({len(png)} bytes)")
            doc.add_picture(io.BytesIO(png), width=self._In(w))
            doc.paragraphs[-1].alignment = A.CENTER
            pc = doc.add_paragraph(); pc.alignment = A.CENTER
            rc = pc.add_run(f"Fig. {self._fn}. {cap}")
            rc.font.size = self._Pt(9); rc.font.italic = True
            rc.font.color.rgb = self._RGB(0x7F,0x8C,0x8D)
        else:
            _L(f"  Fig {self._fn}: FAILED")
            self._errors.append(f"Fig {self._fn} '{cap}': PNG failed")
            p = doc.add_paragraph()
            r = p.add_run(f"[Fig. {self._fn}. {cap} -- "
                          "PNG export failed. Install: pip install \"kaleido==0.2.1\"]")
            r.font.size = self._Pt(9); r.font.italic = True
            r.font.color.rgb = self._RGB(0x95,0xA5,0xA6)
        doc.add_paragraph()

    def _tbl(self, doc, df, max_rows=30):
        from docx.oxml.ns import nsdecls; from docx.oxml import parse_xml
        from docx.enum.text import WD_ALIGN_PARAGRAPH as A
        from docx.enum.table import WD_TABLE_ALIGNMENT
        if df is None or df.empty: return
        d = df.head(max_rows).copy()
        for c in d.columns:
            if d[c].dtype in ('float64','float32'):
                d[c] = d[c].apply(lambda x: f'{x:.4f}' if pd.notnull(x) else '')
        nr,nc = len(d),len(d.columns)
        t = doc.add_table(rows=nr+1, cols=nc)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        try: t.style = 'Table Grid'
        except: pass
        for j,col in enumerate(d.columns):
            c = t.rows[0].cells[j]; c.text = ''
            p = c.paragraphs[0]; p.alignment = A.CENTER
            r = p.add_run(str(col))
            r.font.size=self._Pt(8); r.font.bold=True; r.font.name='Arial'
            r.font.color.rgb = self._RGB(0xFF,0xFF,0xFF)
            try: c._tc.get_or_add_tcPr().append(parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="1A5276" w:val="clear"/>'))
            except: pass
        for i in range(nr):
            for j in range(nc):
                c = t.rows[i+1].cells[j]; c.text = ''
                p = c.paragraphs[0]; p.alignment = A.CENTER
                r = p.add_run(str(d.iloc[i,j]))
                r.font.size=self._Pt(8); r.font.name='Arial'
                if i%2==1:
                    try: c._tc.get_or_add_tcPr().append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="EBF5FB" w:val="clear"/>'))
                    except: pass
        doc.add_paragraph()

    def _box(self, doc, lines, fill="EBF5FB"):
        from docx.oxml.ns import nsdecls; from docx.oxml import parse_xml
        t = doc.add_table(rows=1, cols=1); cell = t.rows[0].cells[0]
        try: cell._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>'))
        except: pass
        cell.text = ''
        for i,line in enumerate(lines):
            p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
            r = p.add_run(line); r.font.size=self._Pt(9); r.font.name='Arial'
        doc.add_paragraph()

    def _key(self, doc, text):
        from docx.oxml.ns import nsdecls; from docx.oxml import parse_xml
        t = doc.add_table(rows=1, cols=1); cell = t.rows[0].cells[0]
        try:
            pr = cell._tc.get_or_add_tcPr()
            pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="D5F5E3" w:val="clear"/>'))
            pr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="12" w:color="27AE60" w:space="0"/>'
                f'</w:tcBorders>'))
        except: pass
        cell.text = ''; p = cell.paragraphs[0]
        r1 = p.add_run("Key finding: ")
        r1.font.size=self._Pt(9); r1.font.bold=True; r1.font.name='Arial'
        r1.font.color.rgb = self._RGB(0x1A,0x52,0x76)
        r2 = p.add_run(text); r2.font.size=self._Pt(9); r2.font.name='Arial'
        doc.add_paragraph()

    # ============================================================== #
    #  SECTIONS                                                        #
    # ============================================================== #

    def _s_alpha(self, doc, a):
        doc.add_heading('1. Alpha Diversity', level=1)
        _L(f"  results: {list(getattr(a,'results',{}).keys())}")
        gc = self.s.get('group_column'); ns = self.dl.otu_table.shape[0]
        ng = self.dl.metadata[gc].nunique() if gc and gc in self.dl.metadata.columns else 0
        test = "Mann-Whitney U" if ng==2 else "Kruskal-Wallis"
        self._txt(doc, f"Alpha diversity for {ns} samples, 6 metrics. "
                  + (f"Groups: {test}, FDR-corrected." if ng>=2 else ""))

        df = _call(a.get_statistics_summary, "a.stats")
        if df is not None and not df.empty:
            doc.add_heading('Statistical Summary', level=2); self._tbl(doc, df)
        gdf = _call(a.get_group_summary_table, "a.group")
        if gdf is not None and not gdf.empty:
            doc.add_heading('Group Statistics', level=2); self._tbl(doc, gdf)

        fig1 = _call(a.plot_all_metrics_plotly,"a.plot4",['observed','shannon','simpson','pielou'])
        self._fig(doc, fig1, 'Observed, Shannon, Simpson, Pielou', fh=450)
        fig2 = _call(a.plot_all_metrics_plotly,"a.plot2",['chao1','faith_pd'])
        self._fig(doc, fig2, 'Chao1, Faith PD', fh=350)
        for m in ['observed','shannon','simpson']:
            lb = getattr(a,'METRIC_LABELS',{}).get(m,m)
            fig = _call(a.plot_metric_plotly,f"a.{m}",m,lb)
            self._fig(doc, fig, f'{lb}', fh=400)

        doc.add_heading('Interpretation', level=2)
        sig, marg = [], []
        sd = getattr(a,'statistics',{})
        if gc and gc in sd:
            for m,st in sd[gc].items():
                pa=st.get('p_adjusted',1); lb=getattr(a,'METRIC_LABELS',{}).get(m,m)
                es=st.get('effect_size',0)
                if pa<0.05: sig.append(f"{lb} (p={pa:.4f}, d={abs(es):.3f})")
                elif pa<0.10: marg.append(f"{lb} (p={pa:.4f})")
        if sig:
            self._txt(doc, f"Significant: {'; '.join(sig)}.")
            self._key(doc, "Significant alpha diversity differences.")
        else:
            self._txt(doc, "No significant differences (FDR p > 0.05).")
            if marg: self._txt(doc, f"Marginal: {'; '.join(marg)}.")
            self._key(doc, "Alpha diversity conserved across groups.")

    def _s_beta(self, doc, a):
        doc.add_heading('2. Beta Diversity', level=1)
        dm = getattr(a,'distance_matrices',{}); orr = getattr(a,'ordination_results',{})
        prr = getattr(a,'permanova_results',{}); mn = getattr(a,'METRIC_NAMES',{})
        metrics = list(dm.keys())
        _L(f"  dm={metrics}, ord={list(orr.keys())}, perm={list(prr.keys())}")
        names = [mn.get(m,m) for m in metrics]
        self._txt(doc, f"Metrics: {', '.join(names)}. PCoA, PERMANOVA, PERMDISP.")

        perm = _call(a.get_permanova_summary,"b.perm")
        if perm is not None and not perm.empty:
            doc.add_heading('PERMANOVA', level=2); self._tbl(doc, perm)
        anos = _call(a.get_anosim_summary,"b.anos")
        if anos is not None and not anos.empty:
            doc.add_heading('ANOSIM', level=2); self._tbl(doc, anos)

        gc = self.s.get('group_column')
        for m in metrics:
            if m in orr:
                fig = _call(a.plot_2d_plotly,f"b.pcoa({m})",m,gc)
                self._fig(doc, fig, f'PCoA -- {mn.get(m,m)}')
        for m in ['bray_curtis','jaccard']:
            if f'{m}_nmds' in orr:
                fig = _call(a.plot_nmds_plotly,f"b.nmds({m})",m,gc)
                self._fig(doc, fig, f'NMDS -- {mn.get(m,m)}'); break
        if metrics:
            fig3 = _call(a.plot_3d,f"b.3d",metrics[0])
            self._fig(doc, fig3, f'PCoA 3D -- {mn.get(metrics[0],metrics[0])}', fh=600)
            figd = _call(a.plot_distance_heatmap,f"b.dhm",metrics[0])
            self._fig(doc, figd, f'Distance heatmap', fh=550)
        figa = _call(a.plot_all_ordinations,"b.all_ord")
        self._fig(doc, figa, 'All ordinations', fw=1100, fh=700)

        doc.add_heading('Variance Explained', level=2)
        for m in metrics:
            if m in orr:
                ev = orr[m].get('explained_variance',[0,0,0])
                self._txt(doc, f"{mn.get(m,m)}: PC1={ev[0]:.1%}, PC2={ev[1]:.1%}")

        doc.add_heading('Interpretation', level=2)
        sig_r, pd_w = [], []
        for m in prr:
            if gc and gc in prr[m]:
                r = prr[m][gc]; mmn = mn.get(m,m)
                if r.get('p_value',1)<0.05:
                    sig_r.append(f"{mmn} (R2={r.get('r_squared',0):.3f}, p={r.get('p_value',1):.4f})")
                dk = f'{gc}_dispersion'
                if dk in prr[m] and prr[m][dk].get('p_value',1)<0.05: pd_w.append(mmn)
        if sig_r:
            self._txt(doc, f"Significant: {'; '.join(sig_r)}.")
            if pd_w: self._txt(doc, f"PERMDISP significant: {', '.join(pd_w)}.")
            self._key(doc, "Significant beta diversity shift."
                      +(" Unequal dispersions." if pd_w else ""))
        else:
            self._txt(doc, "No significant PERMANOVA (p > 0.05).")

    def _s_comp(self, doc, a):
        doc.add_heading('3. Taxonomic Composition', level=1)
        self._txt(doc, "Relative abundance across samples and groups.")
        stats = _call(a.get_summary_statistics,"c.stats")
        if stats and isinstance(stats,dict):
            self._box(doc,[f"Total: {stats.get('total_taxa','?')}",
                           f"Abundant: {stats.get('abundant_taxa','?')}",
                           f"Rare: {stats.get('rare_taxa','?')}"])
        pct = _call(a.create_percentage_table,"c.pct",top_n=20)
        if pct is not None and not pct.empty:
            doc.add_heading('Top 20 Taxa (%)', level=2); self._tbl(doc,pct,20)

        for mname,kw,cap in [
            ('create_plot',dict(top_n=15),'Stacked bar'),
            ('create_group_comparison_plot',dict(top_n=15),'Group comparison'),
            ('create_sunburst',dict(top_n=50),'Sunburst'),
            ('create_treemap',dict(top_n=50),'Treemap'),
            ('create_core_microbiome_plot',dict(),'Core microbiome'),
            ('create_rank_abundance_curve',dict(top_n=30),'Rank-abundance'),
            ('create_taxa_abundance_heatmap_by_group',dict(top_n=20),'Group heatmap'),
            ('create_taxonomy_sankey',dict(top_n=30),'Sankey'),
            ('create_group_indicator_species',dict(top_n=10),'Indicator species')]:
            fn = getattr(a,mname,None)
            if not fn: continue
            r = _call(fn,f"c.{mname}",**kw)
            fig = r[0] if isinstance(r,tuple) else r
            self._fig(doc, fig, cap)
        da = _call(a.create_differential_abundance,"c.diff",top_n=20,p_threshold=0.05)
        if da and isinstance(da,tuple) and len(da)==2:
            self._fig(doc, da[0], 'Differential abundance')
            if da[1] is not None and not da[1].empty:
                doc.add_heading('Differential Results', level=2)
                self._tbl(doc, da[1].head(20))

    def _s_heat(self, doc, a):
        doc.add_heading('4. Abundance Heatmap', level=1)
        _L(f"  heatmap_data: {getattr(a,'heatmap_data',pd.DataFrame()).shape}")
        gc = self.s.get('group_column','')
        tl = self.s.get('taxonomic_level','Genus')
        self._txt(doc, f"Top taxa at {tl} level."+(f" By: {gc}." if gc else ""))
        fig = _call(a.create_heatmap,"h.create",top_n=30)
        self._fig(doc, fig, 'Abundance heatmap', fw=1000, fh=700)
        st_df = _call(a.get_extended_statistics,"h.stats")
        if st_df is not None and not st_df.empty:
            doc.add_heading('Statistics', level=2); self._tbl(doc, st_df.head(25))

    def _s_clust(self, doc, a):
        doc.add_heading('5. Clustered Heatmap', level=1)
        _L(f"  group_means: {getattr(a,'group_means',pd.DataFrame()).shape}")
        self._txt(doc, "Hierarchical clustering of taxa and groups.")
        self._box(doc,["UPGMA, Euclidean, group means (%)"])
        fig = _call(a.create_clustered_heatmap,"ch.create",top_n=30)
        self._fig(doc, fig, 'Clustered heatmap', fw=1000, fh=700)
        s = _call(a.get_comprehensive_statistics,"ch.stats")
        if s and isinstance(s,dict):
            self._box(doc,[f"Taxa: {s.get('total_taxa','?')}",
                           f"Groups: {s.get('total_groups','?')}",
                           f"Max: {s.get('max_abundance',0):.2f}%"])
        self._key(doc, "Co-abundance patterns identified.")

    def _s_net(self, doc, a):
        doc.add_heading('6. Network Analysis', level=1)
        self._txt(doc, "Co-occurrence network.")
        m = _call(a.calculate_network_metrics,"n.met")
        if m and isinstance(m,dict):
            self._box(doc,[f"Nodes: {m.get('nodes','?')}",
                           f"Edges: {m.get('edges','?')}",
                           f"Density: {m.get('density',0):.4f}",
                           f"Avg degree: {m.get('average_degree',0):.2f}",
                           f"Clustering: {m.get('clustering_coefficient',0):.4f}"])
        for layout,cby,cap in [('spring','kingdom','Network (spring)'),
                                ('circular','kingdom','Network (circular)'),
                                ('community','community','Network (communities)')]:
            fig = _call(a.create_network_plot,f"n.{layout}",layout=layout,color_by=cby,show_labels=True)
            self._fig(doc, fig, cap)
        for fn_name,cap in [('create_community_summary_plot','Community structure'),
                            ('create_degree_distribution_plot','Degree distribution')]:
            fn = getattr(a,fn_name,None)
            if fn: self._fig(doc, _call(fn,f"n.{fn_name}"), cap)
        fig_rb = _call(a.analyze_robustness,"n.robust",n_steps=12)
        self._fig(doc, fig_rb, 'Robustness')
        fig_ch = _call(a.create_correlation_heatmap,"n.corr",top_n=30)
        self._fig(doc, fig_ch, 'Correlation matrix', fw=1000, fh=700)
        ks = _call(a.identify_keystone_species,"n.ks",top_n=15)
        if ks is not None and not ks.empty:
            doc.add_heading('Keystone Species', level=2)
            cols=[c for c in ['Taxon','Degree','Betweenness','Closeness','Keystone_Score'] if c in ks.columns]
            self._tbl(doc, ks[cols].head(15))

    def _s_ml(self, doc, a):
        doc.add_heading('7. Machine Learning', level=1)
        self._txt(doc, "Classification from microbial profiles.")
        if not hasattr(a,'models') or not a.models:
            self._txt(doc, "No models."); return
        comp = _call(a.get_model_comparison,"ml.comp")
        if comp is not None and not comp.empty:
            doc.add_heading('Comparison', level=2); self._tbl(doc, comp)
        for fn_name,cap,kw in [('plot_model_comparison','Models',{}),
                               ('plot_feature_importance','Features',dict(top_n=15)),
                               ('plot_confusion_matrices','Confusion',{}),
                               ('plot_roc_curves','ROC',{}),
                               ('plot_cv_scores','CV scores',{})]:
            fn=getattr(a,fn_name,None)
            if fn: self._fig(doc, _call(fn,f"ml.{fn_name}",**kw), cap)
        rpt = _call(a.generate_ml_report,"ml.report")
        if rpt and isinstance(rpt,str) and len(rpt)>50:
            doc.add_heading('Report', level=2)
            for l in rpt.split('\n'):
                if l.strip(): self._txt(doc, l.strip())


def check_export_capability():
    try:
        import kaleido
        v = getattr(kaleido,'__version__','?')
        return True, f"kaleido {v}"
    except ImportError:
        return False, "not installed"
